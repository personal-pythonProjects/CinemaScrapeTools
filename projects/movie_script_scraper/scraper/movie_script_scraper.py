import requests
from bs4 import BeautifulSoup
import os
from docx import Document
from projects.movie_script_scraper.scraper.base_scraper import BaseScraper
from projects.movie_script_scraper.utils.settings import BASE_URL, ENDPOINT, OUTPUT_FOLDER


class MovieScriptScraper(BaseScraper):

    def __init__(self, base_url = BASE_URL, endpoint = ENDPOINT, output_folder = OUTPUT_FOLDER):
        super().__init__(base_url, endpoint, output_folder)

    def fetch_data(self, page):
        url = f"{self.base_url}{self.endpoint}?page={page}"
        response = requests.get(url)

        if response.status_code == 200:
            return response.content
        else:
            raise Exception(f"Failed to fetch data from page {page}")

    def export_to_docx(self, data):

        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)

        for movie in data:

            # Initialize the Document() as an object
            doc = Document()

            # Extract data
            title = movie["title"]
            year = movie["year"]
            plot = movie["plot"]
            transcript = movie["transcript"]
            source = movie["source"]
            link = movie["link"]


            doc.add_heading(f"{title} - {year}", level = 1)
            doc.add_heading(f"Visit:", level = 2)
            doc.add_paragraph(link)
            doc.add_heading("Source:", level = 2)
            doc.add_paragraph(source)

            doc.add_heading("Plot:", level = 2)
            doc.add_paragraph(plot)

            doc.add_heading("Transcript:", level = 2)
            doc.add_paragraph(transcript)

            try:
                filename = os.path.join(self.output_folder, f"{title.title()}.docx")
                doc.save(filename)
            except Exception as e:
                return f"Failed to save file with named {title.title()}.docx. Error: {e}"

        return f"Export the data to docx was successful with {len(data)} entries."

    def parse_data(self, raw_html):
        if raw_html:
            soup = BeautifulSoup(raw_html, 'lxml')
            movie_list = soup.select("ul.scripts-list li a")

            movie_transcripts = []

            for movie in movie_list:
                links = f"{self.base_url}{movie['href']}"

                if links:
                        response = requests.get(links)
                        if response.status_code == 200:
                            soup = BeautifulSoup(response.content, 'lxml')

                            # Extract title
                            data = soup.find("h1").text.split("(")
                            title = data[0].strip().upper() if data[0].strip() else "No title"
                            year = data[1].split(")")[0].strip() if data[1].split(")")[0].strip() else "No year"

                            # Extract plot
                            data = soup.find("p", class_= "plot")
                            plot = data.text.strip() if data else "No plot"

                            # Extract transcript
                            data = soup.find("div", class_= "full-script")
                            transcript = data.text.replace("\n", " ").strip() if data else "No transcript"

                            # Extract source
                            data = soup.select(".page-wrapper span")
                            source = data[0].text if data else "No source"
                            # print(source)

                            movie_transcripts.append({"title": title,
                                                      "year": year,
                                                      "plot": plot,
                                                      "transcript": transcript,
                                                      "source": source,
                                                      "link": links})
                        else:
                            raise Exception(f"Failed to fetch data from {links}")

            result = self.export_to_docx(movie_transcripts)
            return result

        else:
            raise Exception(f"Failed to parse data")

